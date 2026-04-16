from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from PIL import Image
import fitz
import os
import io
import sys
import zipfile
import subprocess
from docx import Document
from docx.shared import Pt
from app.utils.file_helpers import create_temp_file, cleanup_files, get_safe_filename

router = APIRouter()

MAX_FILE_SIZE = 50 * 1024 * 1024

def check_file_size(content: bytes):
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50MB.")

def get_libreoffice_path():
    if sys.platform == "win32":
        return r"C:\Program Files\LibreOffice\program\soffice.exe"
    return "soffice"


@router.post("/pdf-to-images")
async def pdf_to_images(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    input_path = create_temp_file(".pdf")

    try:
        content = await file.read()
        check_file_size(content)
        with open(input_path, "wb") as f:
            f.write(content)

        doc = fitz.open(input_path)
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for i, page in enumerate(doc):
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                zip_file.writestr(f"page_{i + 1}.png", img_bytes)

        doc.close()
        zip_buffer.seek(0)
        filename = f"{get_safe_filename(file.filename)}_pages.zip"

        background_tasks.add_task(cleanup_files, input_path)

        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        cleanup_files(input_path)
        raise
    except Exception as e:
        cleanup_files(input_path)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/word-to-pdf")
async def word_to_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    input_path = create_temp_file(".docx")
    html_path = create_temp_file(".html")
    output_path = create_temp_file(".pdf")

    try:
        content = await file.read()
        check_file_size(content)
        with open(input_path, "wb") as f:
            f.write(content)

        doc = Document(input_path)

        html_parts = ['<!DOCTYPE html><html><head><meta charset="utf-8">']
        html_parts.append('<style>')
        html_parts.append('body { font-family: Arial, sans-serif; margin: 2cm; font-size: 12pt; line-height: 1.5; }')
        html_parts.append('h1 { font-size: 18pt; font-weight: bold; margin: 12pt 0; }')
        html_parts.append('h2 { font-size: 16pt; font-weight: bold; margin: 10pt 0; }')
        html_parts.append('h3 { font-size: 14pt; font-weight: bold; margin: 8pt 0; }')
        html_parts.append('p { margin: 6pt 0; }')
        html_parts.append('table { border-collapse: collapse; width: 100%; }')
        html_parts.append('td, th { border: 1px solid #ccc; padding: 4pt 8pt; }')
        html_parts.append('</style></head><body>')

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                html_parts.append('<br>')
                continue
            style = para.style.name.lower()
            if 'heading 1' in style:
                html_parts.append(f'<h1>{text}</h1>')
            elif 'heading 2' in style:
                html_parts.append(f'<h2>{text}</h2>')
            elif 'heading 3' in style:
                html_parts.append(f'<h3>{text}</h3>')
            else:
                bold = any(run.bold for run in para.runs if run.text.strip())
                if bold:
                    html_parts.append(f'<p><strong>{text}</strong></p>')
                else:
                    html_parts.append(f'<p>{text}</p>')

        for table in doc.tables:
            html_parts.append('<table>')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')

        html_parts.append('</body></html>')
        html_content = ''.join(html_parts)

        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        from weasyprint import HTML
        HTML(filename=html_path).write_pdf(output_path)

        filename = f"{get_safe_filename(file.filename)}.pdf"
        background_tasks.add_task(cleanup_files, input_path, html_path, output_path)

        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename=filename
        )
    except HTTPException:
        cleanup_files(input_path, html_path, output_path)
        raise
    except Exception as e:
        cleanup_files(input_path, html_path, output_path)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pdf-to-word")
async def pdf_to_word(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    input_path = create_temp_file(".pdf")
    output_path = create_temp_file(".docx")

    try:
        content = await file.read()
        check_file_size(content)
        with open(input_path, "wb") as f:
            f.write(content)

        doc = fitz.open(input_path)
        word_doc = Document()

        for page_num, page in enumerate(doc):
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))
            for block in blocks:
                text = block[4].strip()
                if text:
                    para = word_doc.add_paragraph(text)
                    para.style.font.size = Pt(11)
            if page_num < len(doc) - 1:
                word_doc.add_page_break()

        doc.close()
        word_doc.save(output_path)

        filename = f"{get_safe_filename(file.filename)}.docx"

        background_tasks.add_task(cleanup_files, input_path, output_path)

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
    except HTTPException:
        cleanup_files(input_path, output_path)
        raise
    except Exception as e:
        cleanup_files(input_path, output_path)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rotate-pdf")
async def rotate_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    rotation: int = Form(90)
):
    if rotation not in [90, 180, 270]:
        raise HTTPException(status_code=400, detail="Rotation must be 90, 180 or 270 degrees")

    input_path = create_temp_file(".pdf")
    output_path = create_temp_file(".pdf")

    try:
        content = await file.read()
        check_file_size(content)
        with open(input_path, "wb") as f:
            f.write(content)

        doc = fitz.open(input_path)
        for page in doc:
            page.set_rotation(rotation)

        doc.save(output_path)
        doc.close()

        filename = f"{get_safe_filename(file.filename)}_rotated.pdf"

        background_tasks.add_task(cleanup_files, input_path, output_path)

        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename=filename
        )
    except HTTPException:
        cleanup_files(input_path, output_path)
        raise
    except Exception as e:
        cleanup_files(input_path, output_path)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/resize-image")
async def resize_image(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    width: int = Form(None),
    height: int = Form(None)
):
    if not width and not height:
        raise HTTPException(status_code=400, detail="Please provide at least a width or height")

    ext = os.path.splitext(file.filename)[1].lower()
    input_path = create_temp_file(ext)
    output_path = create_temp_file(ext)

    try:
        content = await file.read()
        check_file_size(content)
        with open(input_path, "wb") as f:
            f.write(content)

        img = Image.open(input_path)
        orig_width, orig_height = img.size

        if width and not height:
            ratio = width / orig_width
            height = int(orig_height * ratio)
        elif height and not width:
            ratio = height / orig_height
            width = int(orig_width * ratio)

        resized = img.resize((width, height), Image.LANCZOS)

        save_format = "JPEG" if ext in [".jpg", ".jpeg"] else ext[1:].upper()
        if resized.mode in ["RGBA", "P"] and save_format == "JPEG":
            resized = resized.convert("RGB")

        resized.save(output_path, format=save_format)

        filename = f"{get_safe_filename(file.filename)}_resized{ext}"

        background_tasks.add_task(cleanup_files, input_path, output_path)

        return FileResponse(
            output_path,
            media_type=f"image/{ext[1:]}",
            filename=filename
        )
    except HTTPException:
        cleanup_files(input_path, output_path)
        raise
    except Exception as e:
        cleanup_files(input_path, output_path)
        raise HTTPException(status_code=500, detail=str(e))